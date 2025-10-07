"""Document processor for extracting content from various file formats."""

from __future__ import annotations

import io
from typing import Dict, Any, Optional
import structlog

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """
    Document processor for extracting text and metadata from documents.

    This is a simplified implementation that works with PDF and text files.
    """

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        tenant_id: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        Process document and extract content.

        Args:
            file_content: Raw file content
            filename: Original filename
            tenant_id: Tenant identifier
            **kwargs: Additional processing options

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Determine file type from extension
            file_extension = self._get_file_extension(filename)

            if file_extension == '.pdf':
                return await self._process_pdf(file_content, filename)
            elif file_extension in ['.txt', '.md']:
                return await self._process_text(file_content, filename)
            elif file_extension in ['.doc', '.docx']:
                return await self._process_word(file_content, filename)
            else:
                # Try to process as text for unknown types
                return await self._process_text(file_content, filename)

        except Exception as e:
            logger.error(
                "Error processing document",
                filename=filename,
                error=str(e),
                tenant_id=tenant_id
            )
            raise

    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        if '.' in filename:
            return '.' + filename.split('.')[-1].lower()
        return ''

    async def _process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF file."""
        try:
            # Try to import PyPDF2 or pdfplumber for PDF processing
            import PyPDF2

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""

            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return {
                "text": text.strip(),
                "metadata": {
                    "filename": filename,
                    "file_type": "pdf",
                    "num_pages": len(pdf_reader.pages),
                    "size_bytes": len(content),
                },
            }

        except ImportError:
            logger.warning("PyPDF2 not available, using fallback text extraction")
            return await self._fallback_text_extraction(content, filename, "pdf")
        except Exception as e:
            logger.warning(f"PDF processing failed: {e}, using fallback")
            return await self._fallback_text_extraction(content, filename, "pdf")

    async def _process_text(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process text file."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, use utf-8 with error handling
                text = content.decode('utf-8', errors='replace')

            return {
                "text": text.strip(),
                "metadata": {
                    "filename": filename,
                    "file_type": "text",
                    "size_bytes": len(content),
                    "encoding": encoding,
                },
            }

        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise

    async def _process_word(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word document."""
        try:
            # Try to import python-docx for Word processing
            import docx

            doc = docx.Document(io.BytesIO(content))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return {
                "text": text.strip(),
                "metadata": {
                    "filename": filename,
                    "file_type": "word",
                    "num_paragraphs": len(doc.paragraphs),
                    "size_bytes": len(content),
                },
            }

        except ImportError:
            logger.warning("python-docx not available, using fallback text extraction")
            return await self._fallback_text_extraction(content, filename, "word")
        except Exception as e:
            logger.warning(f"Word processing failed: {e}, using fallback")
            return await self._fallback_text_extraction(content, filename, "word")

    async def _fallback_text_extraction(
        self,
        content: bytes,
        filename: str,
        file_type: str
    ) -> Dict[str, Any]:
        """Fallback text extraction for when proper parsers aren't available."""
        try:
            # Try to extract any readable text
            text = content.decode('utf-8', errors='ignore')

            # Clean up the text (remove null bytes, control characters, etc.)
            text = ''.join(char for char in text if char.isprintable() or char.isspace())

            return {
                "text": text.strip(),
                "metadata": {
                    "filename": filename,
                    "file_type": file_type,
                    "size_bytes": len(content),
                    "extraction_method": "fallback",
                    "warning": f"Used fallback extraction for {file_type} file",
                },
            }

        except Exception as e:
            logger.error(f"Fallback text extraction failed: {e}")
            return {
                "text": "",
                "metadata": {
                    "filename": filename,
                    "file_type": file_type,
                    "size_bytes": len(content),
                    "extraction_method": "failed",
                    "error": str(e),
                },
            }


__all__ = [
    "DocumentProcessor",
]